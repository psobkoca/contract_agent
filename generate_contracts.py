import os
from docx import Document
from fpdf import FPDF
import docx.shared

# Contract templates and details generator
CONTRACT_TYPES = {
    "NDA": {
        "title": "MUTUAL NON-DISCLOSURE AGREEMENT",
        "intro": "This Mutual Non-Disclosure Agreement (the \"Agreement\") is entered into as of {date} (the \"Effective Date\") by and between {party_a} (\"{party_a_short}\") and {party_b} (\"{party_b_short}\").",
        "clauses": [
            ("1. Purpose", "The parties wish to explore a business relationship of mutual interest (the \"Purpose\"). In connection with the Purpose, each party may disclose to the other party certain proprietary, technical, financial, or business information that is confidential in nature."),
            ("2. Definition of Confidential Information", "\"Confidential Information\" means any information disclosed by one party (\"Disclosing Party\") to the other party (\"Receiving Party\") that is marked as confidential or would reasonably be understood to be confidential under the circumstances of disclosure, including but not limited to trade secrets, source code, product plans, and customer lists."),
            ("3. Obligations of Confidentiality", "The Receiving Party agrees to: (a) hold the Confidential Information in strict confidence and use at least the same degree of care to protect it as it uses for its own confidential information; (b) use the Confidential Information solely for the Purpose; and (c) restrict access to Confidential Information to employees and contractors who need to know and are bound by confidentiality obligations at least as restrictive as those herein."),
            ("4. Exceptions", "Confidential Information does not include information that: (a) is or becomes publicly known through no breach of this Agreement; (b) was already in the Receiving Party's possession prior to disclosure; (c) is independently developed by the Receiving Party without reference to or reliance upon the Disclosing Party's Confidential Information; or (d) is rightfully obtained from a third party without restriction."),
            ("5. Term and Termination", "This Agreement and the Receiving Party's obligation to protect Confidential Information shall remain in effect for a period of {term_years} years from the Effective Date, unless terminated earlier by mutual written agreement of the parties."),
            ("6. Governing Law and Jurisdiction", "This Agreement shall be governed by, and construed in accordance with, the laws of the State of {governing_law}, without regard to its principles of conflicts of law. Any legal action arising hereunder shall be brought exclusively in the state or federal courts located in {governing_law}."),
            ("7. Remedies", "The Receiving Party acknowledges that any breach of this Agreement may cause irreparable harm for which monetary damages alone would be inadequate, and that the Disclosing Party shall be entitled to seek injunctive relief in addition to any other remedies available at law or in equity.")
        ]
    },
    "MSA": {
        "title": "MASTER SERVICES AGREEMENT",
        "intro": "This Master Services Agreement (the \"Agreement\") is entered into as of {date} (the \"Effective Date\") by and between {party_a} (\"{party_a_short}\") and {party_b} (\"{party_b_short}\").",
        "clauses": [
            ("1. Scope of Services", "Service Provider shall perform the services (\"Services\") described in one or more Statements of Work (\"SOW\") signed by both parties. Each SOW shall reference this Agreement and be subject to all terms and conditions contained herein."),
            ("2. Fees and Payment Terms", "Client shall pay Service Provider the fees specified in each applicable SOW. Unless otherwise stated in an SOW, all invoices shall be paid within thirty (30) days of the invoice date. Late payments shall accrue interest at a rate of 1.5% per month or the maximum rate permitted by law."),
            ("3. Intellectual Property Ownership", "Except as otherwise explicitly agreed in an SOW, all deliverables, reports, software, and other materials developed by Service Provider in the course of performing the Services shall be the sole and exclusive property of Client, provided that Service Provider retains ownership of its pre-existing intellectual property."),
            ("4. Warranties", "Service Provider warrants that the Services shall be performed in a professional and workmanlike manner in accordance with prevailing industry standards. CLIENT'S SOLE AND EXCLUSIVE REMEDY FOR BREACH OF THIS WARRANTY SHALL BE THE RE-PERFORMANCE OF THE DEFECTIVE SERVICES BY SERVICE PROVIDER."),
            ("5. Limitation of Liability", "IN NO EVENT SHALL EITHER PARTY BE LIABLE TO THE OTHER FOR ANY INDIRECT, INCIDENTAL, SPECIAL, OR CONSEQUENTIAL DAMAGES. SERVICE PROVIDER'S TOTAL CUMULATIVE LIABILITY UNDER THIS AGREEMENT SHALL NOT EXCEED THE TOTAL FEES PAID BY CLIENT TO SERVICE PROVIDER IN THE TWELVE (12) MONTHS PRECEDING THE CLAIM."),
            ("6. Term and Termination", "This Agreement shall commence on the Effective Date and continue until terminated by either party upon thirty (30) days' prior written notice, provided that any active SOW shall survive termination of this Agreement until completed or separately terminated."),
            ("7. Governing Law", "This Agreement and all disputes arising out of or relating to it shall be governed by, and construed in accordance with, the laws of the State of {governing_law}, excluding its conflict of laws rules.")
        ]
    },
    "SaaS": {
        "title": "SAAS SUBSCRIPTION AGREEMENT",
        "intro": "This Software-as-a-Service Subscription Agreement (the \"Agreement\") is entered into as of {date} (the \"Effective Date\") by and between {party_a} (\"{party_a_short}\") and {party_b} (\"{party_b_short}\").",
        "clauses": [
            ("1. SaaS License Grant", "Provider hereby grants Customer a non-exclusive, non-transferable, revocable license to access and use the software-as-a-service platform (the \"Service\") solely for Customer's internal business operations during the subscription term."),
            ("2. Usage Restrictions", "Customer shall not, and shall not permit any third party to: (a) reverse engineer, decompile, or disassemble the Service; (b) copy, modify, or create derivative works of the Service; or (c) bypass or breach any security device or protection used by the Service."),
            ("3. Fees and Billing", "Customer shall pay all fees specified in the Order Form. Subscription fees are billed in advance on a {billing_frequency} basis and are non-refundable. Overdue payments shall be subject to interest charges and may result in the suspension of access to the Service."),
            ("4. Data Security and Privacy", "Provider shall maintain reasonable administrative, physical, and technical safeguards designed to protect the security, confidentiality, and integrity of Customer's data. Provider's collection and use of personal data shall be governed by its Privacy Policy."),
            ("5. Intellectual Property", "As between the parties, Provider retains all right, title, and interest in and to the Service, including all intellectual property rights therein. Customer retains all right, title, and interest in and to all data uploaded by Customer to the Service."),
            ("6. Limitation of Liability", "PROVIDER'S MAXIMUM AGGREGATE LIABILITY ARISING OUT OF OR RELATED TO THIS AGREEMENT, WHETHER IN CONTRACT, TORT, OR OTHERWISE, SHALL NOT EXCEED THE TOTAL FEES PAID BY CUSTOMER TO PROVIDER IN THE TWELVE (12) MONTHS PRECEDING THE INCIDENT GIVING RISE TO LIABILITY."),
            ("7. Governing Law", "This Agreement shall be governed by the laws of the State of {governing_law}, without reference to conflict of laws principles. The parties consent to the exclusive jurisdiction of the courts of {governing_law} for any disputes.")
        ]
    },
    "Vendor": {
        "title": "VENDOR SUPPLY AGREEMENT",
        "intro": "This Vendor Supply Agreement (the \"Agreement\") is entered into as of {date} (the \"Effective Date\") by and between {party_a} (\"{party_a_short}\") and {party_b} (\"{party_b_short}\").",
        "clauses": [
            ("1. Supply of Products", "Vendor agrees to sell and supply to Buyer, and Buyer agrees to purchase, the products (\"Products\") listed in Exhibit A. All purchases shall be initiated by Buyer's issuance of written Purchase Orders to Vendor."),
            ("2. Delivery and Acceptance", "Vendor shall deliver the Products to the location specified in the Purchase Order by the delivery date indicated. Buyer shall have ten (10) business days from receipt to inspect the Products and reject any defective or non-conforming items."),
            ("3. Pricing and Payment", "The prices for the Products shall be as set forth in Exhibit A and shall remain fixed for the initial term of this Agreement. Buyer shall pay all undisputed invoice amounts within forty-five (45) days from the date of invoice receipt."),
            ("4. Title and Risk of Loss", "Title to and risk of loss or damage to the Products shall pass from Vendor to Buyer upon delivery of the Products to Buyer's designated shipping location."),
            ("5. Warranties and Compliance", "Vendor warrants that all Products delivered under this Agreement shall: (a) be new and free from defects in material and workmanship; (b) conform to specifications; and (c) be manufactured and supplied in compliance with all applicable laws and regulations."),
            ("6. Indemnification", "Vendor shall defend, indemnify, and hold harmless Buyer and its affiliates from and against any third-party claims, liabilities, or losses arising out of: (a) any defect in the Products; or (b) any infringement of intellectual property rights by the Products."),
            ("7. Governing Law", "This Agreement shall be governed by and construed in accordance with the laws of the State of {governing_law}, without giving effect to any choice of law or conflict of law provision.")
        ]
    },
    "Partnership": {
        "title": "PARTNERSHIP AGREEMENT",
        "intro": "This Partnership Agreement (the \"Agreement\") is entered into as of {date} (the \"Effective Date\") by and between {party_a} (\"{party_a_short}\") and {party_b} (\"{party_b_short}\").",
        "clauses": [
            ("1. Name and Purpose", "The partners hereby agree to form a general partnership under the name of {partnership_name} (the \"Partnership\"). The primary purpose of the Partnership is {purpose}."),
            ("2. Capital Contributions", "Each Partner shall contribute capital to the Partnership as follows: Partner A ({party_a_short}) shall contribute $100,000, and Partner B ({party_b_short}) shall contribute $100,000. Capital contributions shall be deposited into the Partnership's bank account within fifteen (15) days of the Effective Date."),
            ("3. Allocation of Profits and Losses", "The net profits and losses of the Partnership shall be shared equally (50% to each Partner), unless otherwise agreed in writing signed by both Partners. Distributions of profit shall be made at such times as the Partners may mutually determine."),
            ("4. Management and Voting Rights", "Each Partner shall have an equal voice in the management and conduct of the Partnership business. All major business decisions, including borrowing money, entering into leases, or signing contracts exceeding $10,000, shall require the unanimous consent of both Partners."),
            ("5. Dissolution", "The Partnership may be dissolved at any time by the mutual written agreement of both Partners. Upon dissolution, the assets of the Partnership shall be liquidated, debts shall be paid in order of priority, and any remaining balance shall be distributed to the Partners in proportion to their capital accounts."),
            ("6. Dispute Resolution", "Any dispute or claim arising out of or relating to this Agreement shall be resolved first through good-faith mediation. If mediation is unsuccessful within forty-five (45) days, the dispute shall be submitted to binding arbitration under the rules of the American Arbitration Association."),
            ("7. Governing Law", "This Agreement shall be governed by and construed in accordance with the laws of the State of {governing_law}, without regard to its conflicts of laws principles.")
        ]
    }
}

CONTRACT_DETAILS = [
    # 1. NDA (PDF)
    {
        "id": "CTR_001",
        "ext": "pdf",
        "type": "NDA",
        "date": "January 15, 2026",
        "party_a": "Aegis CyberSecurity Ltd.",
        "party_a_short": "Aegis",
        "party_b": "Zephyr Robotics Inc.",
        "party_b_short": "Zephyr",
        "governing_law": "Delaware",
        "term_years": "3"
    },
    # 2. NDA (DOCX)
    {
        "id": "CTR_002",
        "ext": "docx",
        "type": "NDA",
        "date": "February 12, 2026",
        "party_a": "Obsidian BioTech LLC",
        "party_a_short": "Obsidian",
        "party_b": "Aurora Pharmaceuticals Inc.",
        "party_b_short": "Aurora",
        "governing_law": "Massachusetts",
        "term_years": "5"
    },
    # 3. MSA (PDF)
    {
        "id": "CTR_003",
        "ext": "pdf",
        "type": "MSA",
        "date": "March 1, 2026",
        "party_a": "CloudScale Consulting Group LLC",
        "party_a_short": "CloudScale",
        "party_b": "Apex FinTech Solutions Inc.",
        "party_b_short": "Apex",
        "governing_law": "New York"
    },
    # 4. MSA (DOCX)
    {
        "id": "CTR_004",
        "ext": "docx",
        "type": "MSA",
        "date": "March 20, 2026",
        "party_a": "Stellar Creative Agency LLC",
        "party_a_short": "Stellar",
        "party_b": "Vertex E-Commerce Corp",
        "party_b_short": "Vertex",
        "governing_law": "California"
    },
    # 5. SaaS Subscription (PDF)
    {
        "id": "CTR_005",
        "ext": "pdf",
        "type": "SaaS",
        "date": "April 5, 2026",
        "party_a": "LogiChain Software Systems Inc.",
        "party_a_short": "LogiChain",
        "party_b": "Global Freight Logistics Corp",
        "party_b_short": "Global Freight",
        "governing_law": "Texas",
        "billing_frequency": "monthly"
    },
    # 6. SaaS Subscription (DOCX)
    {
        "id": "CTR_006",
        "ext": "docx",
        "type": "SaaS",
        "date": "April 25, 2026",
        "party_a": "HRFlow AI Corp",
        "party_a_short": "HRFlow",
        "party_b": "TalentAcquire International Inc.",
        "party_b_short": "TalentAcquire",
        "governing_law": "Washington",
        "billing_frequency": "annual"
    },
    # 7. Vendor (PDF)
    {
        "id": "CTR_007",
        "ext": "pdf",
        "type": "Vendor",
        "date": "May 8, 2026",
        "party_a": "Titanium Manufacturing Partners LLC",
        "party_a_short": "Titanium Partners",
        "party_b": "HeavyGear Machinery Corp",
        "party_b_short": "HeavyGear",
        "governing_law": "Ohio"
    },
    # 8. Vendor (DOCX)
    {
        "id": "CTR_008",
        "ext": "docx",
        "type": "Vendor",
        "date": "May 18, 2026",
        "party_a": "EcoPack Packaging Supplies LLC",
        "party_a_short": "EcoPack",
        "party_b": "FreshFoods Grocery Chain LLC",
        "party_b_short": "FreshFoods",
        "governing_law": "Oregon"
    },
    # 9. Partnership (PDF)
    {
        "id": "CTR_009",
        "ext": "pdf",
        "type": "Partnership",
        "date": "June 1, 2026",
        "party_a": "Crimson Media Group LLC",
        "party_a_short": "Crimson",
        "party_b": "Vanguard Gaming Labs Inc.",
        "party_b_short": "Vanguard",
        "governing_law": "Nevada",
        "partnership_name": "Crimson Vanguard Studios",
        "purpose": "joint development and marketing of esports digital media content and gaming broadcasts"
    },
    # 10. Partnership (DOCX)
    {
        "id": "CTR_010",
        "ext": "docx",
        "type": "Partnership",
        "date": "June 15, 2026",
        "party_a": "SilverLine Transit Corp",
        "party_a_short": "SilverLine",
        "party_b": "MetroGreen Mobility LLC",
        "party_b_short": "MetroGreen",
        "governing_law": "Illinois",
        "partnership_name": "SilverGreen Urban Transit Ventures",
        "purpose": "joint operation and maintenance of a municipal electric micro-mobility vehicle network"
    }
]

class PDFContract(FPDF):
    def __init__(self, title_text):
        super().__init__()
        self.title_text = title_text
        
    def header(self):
        self.set_font('helvetica', 'B', 14)
        self.cell(0, 10, self.title_text, border=0, align='C', new_x="LMARGIN", new_y="NEXT")
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}} - CONFIDENTIAL', align='C')

def generate_pdf(details, template, filepath):
    pdf = PDFContract(template["title"])
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_font('helvetica', '', 10)
    
    # Intro
    intro_text = template["intro"].format(**details)
    pdf.multi_cell(0, 6, intro_text, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    
    # Clauses
    for title, content in template["clauses"]:
        pdf.set_font('helvetica', 'B', 11)
        pdf.multi_cell(0, 6, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font('helvetica', '', 10)
        formatted_content = content.format(**details)
        pdf.multi_cell(0, 5, formatted_content, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)
        
    # Signature block
    pdf.ln(10)
    pdf.set_font('helvetica', 'B', 10)
    pdf.cell(90, 5, "PARTY A:")
    pdf.cell(90, 5, "PARTY B:")
    pdf.ln(5)
    pdf.set_font('helvetica', '', 10)
    pdf.cell(90, 5, details["party_a"])
    pdf.cell(90, 5, details["party_b"])
    pdf.ln(15)
    pdf.cell(90, 5, "By: ___________________________")
    pdf.cell(90, 5, "By: ___________________________")
    pdf.ln(5)
    pdf.cell(90, 5, "Name: ")
    pdf.cell(90, 5, "Name: ")
    pdf.ln(5)
    pdf.cell(90, 5, "Title: ")
    pdf.cell(90, 5, "Title: ")
    
    pdf.output(filepath)

def generate_docx(details, template, filepath):
    doc = Document()
    
    # Document title
    title = doc.add_paragraph()
    run = title.add_run(template["title"])
    run.bold = True
    run.font.size = docx.shared.Pt(14)
    title.alignment = 1  # Center
    
    # Intro
    intro_text = template["intro"].format(**details)
    doc.add_paragraph(intro_text)
    
    # Clauses
    for title, content in template["clauses"]:
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(title)
        run_title.bold = True
        
        formatted_content = content.format(**details)
        doc.add_paragraph(formatted_content)
        
    # Signature block
    p_sig = doc.add_paragraph()
    p_sig.add_run("\nIN WITNESS WHEREOF, the parties hereto have executed this Agreement as of the Effective Date.\n").italic = True
    
    # Create signature table
    table = doc.add_table(rows=5, cols=2)
    table.autofit = False
    
    # Set headers
    table.cell(0, 0).paragraphs[0].add_run("PARTY A: " + details["party_a"]).bold = True
    table.cell(0, 1).paragraphs[0].add_run("PARTY B: " + details["party_b"]).bold = True
    
    table.cell(2, 0).paragraphs[0].add_run("By: ___________________________")
    table.cell(2, 1).paragraphs[0].add_run("By: ___________________________")
    
    table.cell(3, 0).paragraphs[0].add_run("Name: ")
    table.cell(3, 1).paragraphs[0].add_run("Name: ")
    
    table.cell(4, 0).paragraphs[0].add_run("Title: ")
    table.cell(4, 1).paragraphs[0].add_run("Title: ")
    
    doc.save(filepath)

def main():
    os.makedirs("contracts", exist_ok=True)
    print("Generating contracts in contracts/ directory...")
    
    for details in CONTRACT_DETAILS:
        template = CONTRACT_TYPES[details["type"]]
        filename = f"{details['id']}.{details['ext']}"
        filepath = os.path.join("contracts", filename)
        
        if details["ext"] == "pdf":
            generate_pdf(details, template, filepath)
            print(f"Generated PDF: {filepath}")
        elif details["ext"] == "docx":
            generate_docx(details, template, filepath)
            print(f"Generated DOCX: {filepath}")
            
    print("All 10 contracts successfully generated!")

if __name__ == "__main__":
    main()
